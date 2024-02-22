import openpyxl
from docx import Document
import os
from datetime import datetime
from docx.shared import Pt

def replace_placeholder_with_excel_data(excel_file_path, word_file_path, result_file_path):
    """
    This function will replace placeholders in a Word document with data from an Excel file.
    The placeholders in the Word document should be formatted as {{placeholder1}}, {{placeholder2}}, etc.
    The data will be taken from the last row of the Excel file.
    """
    
    # Load the Excel file
    wb = openpyxl.load_workbook(excel_file_path)
    sheet = wb.active

    # Get the last row in the Excel file
    last_row = sheet.max_row
    data = [cell.value for cell in sheet[last_row]]

    for d in data:
        if isinstance(d, datetime):
            data[data.index(d)] = d.strftime('%d %b %Y')

    # Load the Word document
    doc = Document(word_file_path)

    # Replace placeholders in the Word document with data from the Excel file
    for paragraph in doc.paragraphs:
        # print([run.text for run in paragraph.runs])

        for i in range(len(data)):
            if '{{placeholder' + str(i+1) + '}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{placeholder' + str(i+1) + '}}', str(data[i]))

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Check if the cell has only one paragraph and no runs
                if len(cell.paragraphs) == 1 and len(cell.paragraphs[0].runs) == 0:
                    for i in range(len(data)):
                        if '{{placeholder' + str(i+1) + '}}' in cell.text:
                            cell.text = cell.text.replace('{{placeholder' + str(i+1) + '}}', str(data[i]))
                else:
                    for paragraph in cell.paragraphs:
                        
                        for i in range(len(data)):
                            if '{{placeholder' + str(i+1) + '}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{placeholder' + str(i+1) + '}}', str(data[i]))

    # Save the result
    doc.save(result_file_path)



def replace_placeholder_with_excel_data2(excel_file_path, word_file_path, result_file_path):
    """
    This function will replace placeholders in a Word document with data from an Excel file.
    The placeholders in the Word document should be formatted as {{placeholder1}}, {{placeholder2}}, etc.
    The data will be taken from the last row of the Excel file.
    """
    
    # Load the Excel file
    wb = openpyxl.load_workbook(excel_file_path)
    sheet = wb.active

    # Get the last row in the Excel file
    last_row = sheet.max_row
    data = [cell.value for cell in sheet[last_row]]

    # Load the Word document
    doc = Document(word_file_path)

    # Replace placeholders in the Word document with data from the Excel file
    for paragraph in doc.paragraphs:
        for i in range(len(data)):
            placeholder = '{{placeholder' + str(i+1) + '}}'
            if placeholder in paragraph.text:
                runs = paragraph.runs
                print(runs)
                for run in runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(data[i]))
                        break

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Check if the cell has only one paragraph and no runs
                if len(cell.paragraphs) == 1 and len(cell.paragraphs[0].runs) == 0:
                    for i in range(len(data)):
                        placeholder = '{{placeholder' + str(i+1) + '}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(data[i]))
                else:
                    for paragraph in cell.paragraphs:
                        for i in range(len(data)):
                            placeholder = '{{placeholder' + str(i+1) + '}}'
                            if placeholder in paragraph.text:
                                runs = paragraph.runs
                                for run in runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(data[i]))
                                        break

    # Save the result
    doc.save(result_file_path)


def read_filename(excel_file_path):
    wb = openpyxl.load_workbook(excel_file_path)
    sheet = wb.active
    filename= sheet.cell(row=sheet.max_row,column=14).value
    return filename


def get_text_excluding_headers_footers(word_file_path):
    doc = Document(word_file_path)
    words = []

    # Iterate through each paragraph in the document body
    for para in doc.paragraphs:
        words.extend(para.text.split())

    # Iterate through each table in the document body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    words.extend(para.text.split())

    return words



if __name__ == "__main__":
    # Define the file paths
    # current_dir=os.path.dirname(os.path.abspath(__file__))
    
    excel_file_path = os.path.dirname(os.path.abspath(__file__))+''.join(["\\","input_castin.xlsx"])
    # print(excel_file_path)
    word_file_path = os.path.dirname(os.path.abspath(__file__))+''.join(["\\","template_castin.docx"])
    


    # Get all words from the document content and tables, excluding headers and footers
    words = get_text_excluding_headers_footers(word_file_path)

    # Print the list of words or process them as needed
    # print(words)

    output_filename=read_filename(excel_file_path)
    result_file_path = os.path.dirname(os.path.abspath(__file__))+''.join(["\\",str(output_filename),".docx"])
    # Call the function
    replace_placeholder_with_excel_data(excel_file_path, word_file_path, result_file_path)
    print('_______________________________________________________')